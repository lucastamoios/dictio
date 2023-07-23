# import argparse

from words.models import Word

from alive_progress import alive_bar
from django.core.management.base import BaseCommand, CommandError
from docx import Document

def to_italic(t):
    return f'<i>{t}</i>'


def to_bold(t):
    return f'<b>{t}</b>'

class Command(BaseCommand):
    help = "Imports latin words from a .docx file"

    def add_arguments(self, parser):
        parser.add_argument(
            "file",
            type=ascii,
            help="The file containing the words to import",
        )

    def handle(self, *args, **options):
        try:
            filename = options["file"].strip("'")
            f = open(filename, "rb")
            doc = Document(f)
            table = doc.tables[0]
            with alive_bar(len(table.rows)) as bar:
                for row in table.rows:
                    cells = []
                    for cell in row.cells:
                        text = ''
                        # Docs will have multiple paragraphs in a cell, each with multiple runs, we must join them together
                        # to get the full text, taking care of the formatting.
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                t = run.text
                                if run.bold:
                                    t = to_bold(t)
                                if run.italic:
                                    t = to_italic(t)
                                text += t
                        cells.append(text)
                    # After we have all the data from the row, we can save it to the database
                    Word.objects.create(word=cells[0], meta=cells[1], definition=cells[2])
                    bar()
            
            self.stdout.write(self.style.SUCCESS(f"Finished adding {len(table.rows)} words to the database."))
        except Exception as e:
            self.stdout.write(self.style.ERROR(e))
        finally:
            self.stdout.write(self.style.NOTICE("Closing file."))
            f.close()